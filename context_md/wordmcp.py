"""
Word Document Module

Provides Word document creation using python-docx library for document manipulation.
"""

import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class WordMCPClient:
    """Client for Word document generation using python-docx."""
    
    def __init__(self):
        """Initialize Word document client."""
        self.documents: Dict[str, Document] = {}
        logger.info("Word document client initialized")
    
    def _get_document(self, filename: str) -> Document:
        """Get or create document instance.
        
        Args:
            filename: Path to document file
            
        Returns:
            Document instance
        """
        if filename not in self.documents:
            self.documents[filename] = Document()
        return self.documents[filename]
    
    def create_document(self, filename: str, title: Optional[str] = None, 
                       author: Optional[str] = None) -> str:
        """Create a new Word document.
        
        Args:
            filename: Path to output DOCX file
            title: Document title
            author: Document author
            
        Returns:
            Success message
        """
        doc = Document()
        self.documents[filename] = doc
        
        # Set document properties
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author
        
        logger.info(f"Created document: {filename}")
        return "Document created successfully"
    
    def add_heading(self, filename: str, text: str, level: int = 1,
                   font_name: Optional[str] = None, font_size: Optional[int] = None,
                   bold: Optional[bool] = None, italic: Optional[bool] = None) -> str:
        """Add a heading to the document.
        
        Args:
            filename: Path to DOCX file
            text: Heading text
            level: Heading level (1-9)
            font_name: Font family name (ignored for headings)
            font_size: Font size in points (ignored for headings)
            bold: Bold formatting (ignored for headings)
            italic: Italic formatting (ignored for headings)
            
        Returns:
            Success message
        """
        doc = self._get_document(filename)
        doc.add_heading(text, level=level)
        logger.debug(f"Added heading (level {level}): {text[:50]}...")
        return "Heading added successfully"
    
    def add_paragraph(self, filename: str, text: str, style: Optional[str] = None,
                     font_name: Optional[str] = None, font_size: Optional[int] = None,
                     bold: Optional[bool] = None, italic: Optional[bool] = None,
                     color: Optional[str] = None) -> str:
        """Add a paragraph to the document.
        
        Args:
            filename: Path to DOCX file
            text: Paragraph text
            style: Paragraph style name
            font_name: Font family name
            font_size: Font size in points
            bold: Bold formatting
            italic: Italic formatting
            color: Text color (hex format)
            
        Returns:
            Success message
        """
        doc = self._get_document(filename)
        paragraph = doc.add_paragraph(text, style=style)
        
        # Apply formatting to paragraph runs
        if font_name or font_size or bold or italic or color:
            for run in paragraph.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold:
                    run.font.bold = True
                if italic:
                    run.font.italic = True
                if color:
                    # Parse hex color (e.g., "#FF0000")
                    color = color.lstrip('#')
                    r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
        
        logger.debug(f"Added paragraph: {text[:50]}...")
        return "Paragraph added successfully"
    
    def add_table(self, filename: str, rows: int, cols: int, 
                 data: Optional[List[List[str]]] = None) -> str:
        """Add a table to the document.
        
        Args:
            filename: Path to DOCX file
            rows: Number of rows
            cols: Number of columns
            data: Table data as 2D list
            
        Returns:
            Success message
        """
        doc = self._get_document(filename)
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Light Grid Accent 1'
        
        # Populate table data if provided
        if data:
            for i, row_data in enumerate(data):
                if i < len(table.rows):
                    row = table.rows[i]
                    for j, cell_text in enumerate(row_data):
                        if j < len(row.cells):
                            row.cells[j].text = str(cell_text)
        
        logger.debug(f"Added table: {rows}x{cols}")
        return "Table added successfully"
    
    def add_page_break(self, filename: str) -> str:
        """Add a page break to the document.
        
        Args:
            filename: Path to DOCX file
            
        Returns:
            Success message
        """
        doc = self._get_document(filename)
        doc.add_page_break()
        logger.debug("Added page break")
        return "Page break added successfully"
    
    def save_document(self, filename: str) -> None:
        """Save document to file.
        
        Args:
            filename: Path to output file
        """
        if filename in self.documents:
            self.documents[filename].save(filename)
            logger.info(f"Saved document: {filename}")
    
    def convert_to_pdf(self, filename: str, output_filename: Optional[str] = None) -> str:
        """Convert DOCX to PDF (not supported by python-docx).
        
        Args:
            filename: Path to input DOCX file
            output_filename: Path to output PDF file
            
        Returns:
            Error message indicating PDF conversion requires external tools
        """
        raise NotImplementedError("PDF conversion requires external tools (use _convert_docx_to_pdf)")
    
    def __enter__(self):
        """Context manager entry."""
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        """Context manager exit - save all documents."""
        for filename in list(self.documents.keys()):
            try:
                self.save_document(filename)
            except Exception as e:
                logger.error(f"Failed to save document {filename}: {e}")
        return False
