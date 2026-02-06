"""
Export Command - Convert Markdown files to DOCX and PDF formats

Usage:
    context-md export docs/prd/PRD-123.md --format docx
    context-md export docs/adr/ADR-456.md --format pdf
    context-md export --issue 123 --format both
"""

import json
import logging
import os
import subprocess
from pathlib import Path
from typing import Optional, List

import click

from context_md.state import State
from context_md.wordmcp import WordMCPClient

logger = logging.getLogger(__name__)


@click.group("export")
@click.pass_context
def export_cmd(_ctx: click.Context) -> None:
    """Export markdown documents to DOCX and PDF formats."""
    pass


@export_cmd.command("document")
@click.argument("filepath", type=click.Path(exists=True))
@click.option("--format", "-f", type=click.Choice(["docx", "pdf", "both"]), default="docx",
              help="Output format")
@click.option("--output", "-o", type=click.Path(), help="Output directory (default: same as source)")
@click.option("--template", "-t", type=click.Path(exists=True),
             help="Template file for styling")
@click.pass_context
def export_document_cmd(ctx: click.Context, filepath: str, format: str,
                       output: Optional[str], template: Optional[str]) -> None:
    """Export a markdown document to DOCX/PDF format.
    
    Examples:
        context-md export document docs/prd/PRD-123.md --format docx
        context-md export document docs/adr/ADR-456.md --format pdf --output ./deliverables/
    """
    filepath_obj = Path(filepath)
    
    if not filepath_obj.suffix == ".md":
        raise click.ClickException("Only markdown (.md) files are supported")
    
    # Determine output directory
    output_dir = Path(output) if output else filepath_obj.parent
    output_dir.mkdir(parents=True, exist_ok=True)
    
    base_name = filepath_obj.stem
    
    click.echo(f"[SUCCESS] Exporting {filepath_obj.name}...")
    
    # Convert to DOCX
    if format in ["docx", "both"]:
        docx_path = output_dir / f"{base_name}.docx"
        try:
            _convert_md_to_docx(filepath_obj, docx_path, template)
            click.echo(f"[SUCCESS] DOCX created: {docx_path}")
        except Exception as e:
            click.echo(f"[ERROR] DOCX conversion failed: {e}", err=True)
    
    # Convert to PDF
    if format in ["pdf", "both"]:
        pdf_path = output_dir / f"{base_name}.pdf"
        try:
            if format == "both":
                # Convert DOCX to PDF using Word MCP
                _convert_docx_to_pdf(docx_path, pdf_path)
            else:
                # Direct MD to PDF
                _convert_md_to_pdf(filepath_obj, pdf_path)
            click.echo(f"[SUCCESS] PDF created: {pdf_path}")
        except Exception as e:
            click.echo(f"[ERROR] PDF conversion failed: {e}", err=True)


@export_cmd.command("issue")
@click.argument("issue_number", type=int)
@click.option("--format", "-f", type=click.Choice(["docx", "pdf", "both"]), default="both",
              help="Output format")
@click.option("--output", "-o", type=click.Path(), default="./deliverables",
              help="Output directory")
@click.pass_context
def export_issue_cmd(ctx: click.Context, issue_number: int, format: str,
                    output: str) -> None:
    """Export all deliverables for an issue (PRD, ADR, Spec, Review).
    
    Examples:
        context-md export issue 123 --format both --output ./deliverables/
    """
    repo_root = ctx.obj.get("repo_root")
    if not repo_root:
        raise click.ClickException("Not in a Git repository")
    
    output_dir = Path(output)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Find all documents related to the issue
    docs = _find_issue_documents(repo_root, issue_number)
    
    if not docs:
        click.echo(f"[WARNING] No documents found for issue #{issue_number}")
        return
    
    click.echo(f"[INFO] Found {len(docs)} document(s) for issue #{issue_number}")
    
    for doc in docs:
        click.echo(f"\n[INFO] Processing {doc.name}...")
        
        base_name = doc.stem
        
        # Convert to DOCX
        if format in ["docx", "both"]:
            docx_path = output_dir / f"{base_name}.docx"
            try:
                _convert_md_to_docx(doc, docx_path, None)
                click.echo(f"  [SUCCESS] DOCX: {docx_path}")
            except Exception as e:
                click.echo(f"  [ERROR] DOCX failed: {e}", err=True)
        
        # Convert to PDF
        if format in ["pdf", "both"]:
            pdf_path = output_dir / f"{base_name}.pdf"
            try:
                if format == "both":
                    _convert_docx_to_pdf(docx_path, pdf_path)
                else:
                    _convert_md_to_pdf(doc, pdf_path)
                click.echo(f"  [SUCCESS] PDF: {pdf_path}")
            except Exception as e:
                click.echo(f"  [ERROR] PDF failed: {e}", err=True)
    
    click.echo(f"\n[SUCCESS] Export complete! Files saved to: {output_dir}")


def _extract_table_from_tokens(tokens: List, start_index: int) -> Optional[List[List[str]]]:
    """Extract table data from markdown tokens.
    
    Args:
        tokens: List of markdown tokens
        start_index: Index of table_open token
        
    Returns:
        Table data as 2D list, or None if extraction fails
    """
    table_data = []
    i = start_index + 1
    
    # Find table rows
    while i < len(tokens) and tokens[i].type != "table_close":
        if tokens[i].type == "tr_open":
            row_data = []
            i += 1
            
            # Extract cells in this row
            while i < len(tokens) and tokens[i].type != "tr_close":
                if tokens[i].type in ["th_open", "td_open"]:
                    i += 1
                    if i < len(tokens) and tokens[i].type == "inline":
                        row_data.append(tokens[i].content)
                i += 1
            
            if row_data:
                table_data.append(row_data)
        i += 1
    
    return table_data if table_data else None


def _find_issue_documents(repo_root: Path, issue_number: int) -> List[Path]:
    """Find all deliverable documents for an issue."""
    docs = []
    
    # Search patterns
    patterns = [
        f"docs/prd/PRD-{issue_number}*.md",
        f"docs/adr/ADR-{issue_number}*.md",
        f"docs/specs/SPEC-{issue_number}*.md",
        f"docs/ux/UX-{issue_number}*.md",
        f"docs/reviews/REVIEW-{issue_number}*.md",
    ]
    
    for pattern in patterns:
        docs.extend(repo_root.glob(pattern))
    
    return sorted(docs)


def _convert_md_to_docx(md_path: Path, docx_path: Path, template: Optional[Path]) -> None:
    """Convert markdown to DOCX using Word MCP Server."""
    from markdown_it import MarkdownIt
    
    # Parse markdown content
    md = MarkdownIt("commonmark", {"breaks": True, "html": False})
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()
    
    tokens = md.parse(content)
    
    # Use Word MCP client to create document
    with WordMCPClient() as mcp:
        _create_word_document_via_mcp(mcp, tokens, docx_path, md_path.stem)


def _create_word_document_via_mcp(mcp: WordMCPClient, tokens: List, output_path: Path,
                                   document_title: str) -> None:
    """Create Word document using MCP server tools."""
    filename = str(output_path)
    
    # Create document via MCP
    logger.info(f"Creating Word document via MCP: {filename}")
    mcp.create_document(filename, title=document_title, author="Context.md Export")
    
    # Process markdown tokens and add content via MCP
    i = 0
    while i < len(tokens):
        token = tokens[i]
        
        try:
            if token.type == "heading_open":
                # Get heading level and text
                level = int(token.tag[1])
                i += 1
                if i < len(tokens) and tokens[i].type == "inline":
                    text = tokens[i].content
                    mcp.add_heading(filename, text, level=level)
                    logger.debug(f"Added heading (level {level}): {text[:50]}...")
            
            elif token.type == "paragraph_open":
                # Get paragraph text
                i += 1
                if i < len(tokens) and tokens[i].type == "inline":
                    text = tokens[i].content
                    if text:  # Only add non-empty paragraphs
                        mcp.add_paragraph(filename, text)
                        logger.debug(f"Added paragraph: {text[:50]}...")
            
            elif token.type == "code_block" or token.type == "fence":
                # Add code block as styled paragraph
                code_text = token.content
                if code_text:
                    mcp.add_paragraph(filename, code_text, 
                                    font_name="Courier New",
                                    font_size=10)
                    logger.debug(f"Added code block: {len(code_text)} chars")
            
            elif token.type == "table_open":
                # Handle tables - collect table data
                table_data = _extract_table_from_tokens(tokens, i)
                if table_data:
                    rows = len(table_data)
                    cols = len(table_data[0]) if rows > 0 else 0
                    mcp.add_table(filename, rows, cols, table_data)
                    logger.debug(f"Added table: {rows}x{cols}")
            
            elif token.type == "hr":
                # Add horizontal rule as page break
                mcp.add_page_break(filename)
                logger.debug("Added page break")
        
        except Exception as e:
            logger.warning(f"Failed to add element (token type: {token.type}): {e}")
        
        i += 1
    
    logger.info(f"Document created successfully: {filename}")


def _convert_md_to_pdf(md_path: Path, pdf_path: Path) -> None:
    """Convert markdown to PDF using weasyprint or pandoc."""
    # Try weasyprint first (better CSS support)
    try:
        import markdown
        from weasyprint import HTML
        
        # Convert MD to HTML
        with open(md_path, "r", encoding="utf-8") as f:
            md_content = f.read()
        
        html_content = markdown.markdown(md_content, extensions=[
            "tables", "fenced_code", "codehilite", "toc"
        ])
        
        # Add CSS styling
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    line-height: 1.6;
                    max-width: 800px;
                    margin: 40px auto;
                    padding: 20px;
                }}
                h1, h2, h3 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
                code {{ background: #f4f4f4; padding: 2px 5px; border-radius: 3px; }}
                pre {{ background: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }}
                table {{ border-collapse: collapse; width: 100%; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #3498db; color: white; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # Convert HTML to PDF
        HTML(string=styled_html).write_pdf(str(pdf_path))
        
    except ImportError:
        # Fallback to pandoc
        try:
            subprocess.run([
                "pandoc",
                str(md_path),
                "-o", str(pdf_path),
                "--pdf-engine=xelatex",
                "-V", "geometry:margin=1in"
            ], check=True)
        except FileNotFoundError:
            raise RuntimeError(
                "PDF conversion requires either:\n"
                "  1. pip install weasyprint markdown\n"
                "  2. Install pandoc: https://pandoc.org/installing.html"
            )


def _convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """Convert DOCX to PDF using platform-specific tools.
    
    Args:
        docx_path: Path to input DOCX file
        pdf_path: Path to output PDF file
        
    Raises:
        RuntimeError: If PDF conversion fails
    """
    logger.info(f"Converting DOCX to PDF: {docx_path} -> {pdf_path}")
    
    import platform
    
    try:
        if platform.system() == "Windows":
            # Try docx2pdf (uses Microsoft Word)
            try:
                from docx2pdf import convert
                convert(str(docx_path), str(pdf_path))
                logger.info("PDF created using docx2pdf (Microsoft Word)")
                return
            except ImportError:
                logger.warning("docx2pdf not available, using weasyprint fallback")
                pass
            
            # Fallback: Convert DOCX content to PDF via weasyprint
            try:
                from docx import Document
                from weasyprint import HTML
                
                # Extract text content from DOCX
                doc = Document(str(docx_path))
                paragraphs = []
                for p in doc.paragraphs:
                    if p.text:
                        paragraphs.append(f"<p>{p.text}</p>")
                
                html_content = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="utf-8">
                    <style>
                        body {{
                            font-family: Arial, sans-serif;
                            line-height: 1.6;
                            max-width: 800px;
                            margin: 40px auto;
                            padding: 20px;
                        }}
                        p {{ margin-bottom: 10px; }}
                    </style>
                </head>
                <body>
                    {''.join(paragraphs)}
                </body>
                </html>
                """
                
                HTML(string=html_content).write_pdf(str(pdf_path))
                logger.info("PDF created using weasyprint from DOCX content")
                return
            except Exception as e:
                raise RuntimeError(
                    f"Windows PDF conversion failed: {e}\n"
                    "Install with: pip install docx2pdf\n"
                    "Note: Requires Microsoft Word to be installed"
                )
        else:
            # Try LibreOffice on Linux/Mac
            subprocess.run([
                "soffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(pdf_path.parent),
                str(docx_path)
            ], check=True)
            logger.info("PDF created using LibreOffice")
            
    except Exception as e:
        raise RuntimeError(
            f"PDF conversion failed: {e}\n"
            "Install one of:\n"
            "  1. pip install docx2pdf (Windows, requires MS Word)\n"
            "  2. sudo apt install libreoffice (Linux)\n"
            "  3. brew install libreoffice (Mac)"
        )

