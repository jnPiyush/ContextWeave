"""
Export Command - Convert Markdown files to DOCX and PDF formats

Usage:
    context-weave export docs/prd/PRD-123.md --format docx
    context-weave export docs/adr/ADR-456.md --format pdf
    context-weave export --issue 123 --format both
"""

import logging
import subprocess
from pathlib import Path
from typing import List, Optional

import click

logger = logging.getLogger(__name__)


@click.group("export")
@click.pass_context
def export_cmd(_ctx: click.Context) -> None:
    """Export markdown documents to DOCX and PDF formats."""


@export_cmd.command("document")
@click.argument("filepath", type=click.Path(exists=True))
@click.option("--format", "-f", "output_format", type=click.Choice(["docx", "pdf", "both"]), default="docx",
              help="Output format")
@click.option("--output", "-o", type=click.Path(), help="Output directory (default: same as source)")
@click.option("--template", "-t", type=click.Path(exists=True),
             help="Template file for styling")
@click.pass_context
def export_document_cmd(_ctx: click.Context, filepath: str, output_format: str,
                       output: Optional[str], template: Optional[str]) -> None:
    """Export a markdown document to DOCX/PDF format.

    Examples:
        context-weave export document docs/prd/PRD-123.md --format docx
        context-weave export document docs/adr/ADR-456.md --format pdf --output ./deliverables/
    """
    filepath_obj = Path(filepath)

    if not filepath_obj.suffix == ".md":
        raise click.ClickException("Only markdown (.md) files are supported")

    # Determine output directory
    output_dir = Path(output) if output else filepath_obj.parent
    output_dir.mkdir(parents=True, exist_ok=True)

    base_name = filepath_obj.stem

    click.echo(f"[SUCCESS] Exporting {filepath_obj.name}...")

    # Convert to DOCX
    if output_format in ["docx", "both"]:
        docx_path = output_dir / f"{base_name}.docx"
        try:
            template_path = Path(template) if template else None
            _convert_md_to_docx(filepath_obj, docx_path, template_path)
            click.echo(f"[SUCCESS] DOCX created: {docx_path}")
        except (OSError, RuntimeError, ImportError) as e:
            click.echo(f"[ERROR] DOCX conversion failed: {e}", err=True)

    # Convert to PDF
    if output_format in ["pdf", "both"]:
        pdf_path = output_dir / f"{base_name}.pdf"
        try:
            if output_format == "both":
                # Convert DOCX to PDF using Word MCP
                _convert_docx_to_pdf(docx_path, pdf_path)
            else:
                # Direct MD to PDF
                _convert_md_to_pdf(filepath_obj, pdf_path)
            click.echo(f"[SUCCESS] PDF created: {pdf_path}")
        except (OSError, RuntimeError, ImportError) as e:
            click.echo(f"[ERROR] PDF conversion failed: {e}", err=True)


@export_cmd.command("issue")
@click.argument("issue_number", type=int)
@click.option("--format", "-f", "output_format", type=click.Choice(["docx", "pdf", "both"]), default="both",
              help="Output format")
@click.option("--output", "-o", type=click.Path(), default="./deliverables",
              help="Output directory")
@click.pass_context
def export_issue_cmd(ctx: click.Context, issue_number: int, output_format: str,
                    output: str) -> None:
    """Export all deliverables for an issue (PRD, ADR, Spec, Review).

    Examples:
        context-weave export issue 123 --format both --output ./deliverables/
    """
    repo_root = ctx.obj.get("repo_root")
    if not repo_root:
        raise click.ClickException("Not in a Git repository")

    output_dir = Path(output)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Find all documents related to the issue
    docs = _find_issue_documents(repo_root, issue_number)

    if not docs:
        click.echo(f"[WARNING] No documents found for issue #{issue_number}")
        return

    click.echo(f"[INFO] Found {len(docs)} document(s) for issue #{issue_number}")

    for doc in docs:
        click.echo(f"\n[INFO] Processing {doc.name}...")

        base_name = doc.stem

        # Convert to DOCX
        if output_format in ["docx", "both"]:
            docx_path = output_dir / f"{base_name}.docx"
            try:
                _convert_md_to_docx(doc, docx_path, None)
                click.echo(f"  [SUCCESS] DOCX: {docx_path}")
            except (OSError, RuntimeError, ImportError) as e:
                click.echo(f"  [ERROR] DOCX failed: {e}", err=True)

        # Convert to PDF
        if output_format in ["pdf", "both"]:
            pdf_path = output_dir / f"{base_name}.pdf"
            try:
                if output_format == "both":
                    _convert_docx_to_pdf(docx_path, pdf_path)
                else:
                    _convert_md_to_pdf(doc, pdf_path)
                click.echo(f"  [SUCCESS] PDF: {pdf_path}")
            except (OSError, RuntimeError, ImportError) as e:
                click.echo(f"  [ERROR] PDF failed: {e}", err=True)

    click.echo(f"\n[SUCCESS] Export complete! Files saved to: {output_dir}")


def _extract_table_from_tokens(tokens: List, start_index: int) -> Optional[List[List[str]]]:
    """Extract table data from markdown tokens.

    Args:
        tokens: List of markdown tokens
        start_index: Index of table_open token

    Returns:
        Table data as 2D list, or None if extraction fails
    """
    table_data = []
    i = start_index + 1

    # Find table rows
    while i < len(tokens) and tokens[i].type != "table_close":
        if tokens[i].type == "tr_open":
            row_data = []
            i += 1

            # Extract cells in this row
            while i < len(tokens) and tokens[i].type != "tr_close":
                if tokens[i].type in ["th_open", "td_open"]:
                    i += 1
                    if i < len(tokens) and tokens[i].type == "inline":
                        row_data.append(tokens[i].content)
                i += 1

            if row_data:
                table_data.append(row_data)
        i += 1

    return table_data if table_data else None


def _find_issue_documents(repo_root: Path, issue_number: int) -> List[Path]:
    """Find all deliverable documents for an issue."""
    docs: List[Path] = []

    # Search patterns
    patterns = [
        f"docs/prd/PRD-{issue_number}*.md",
        f"docs/adr/ADR-{issue_number}*.md",
        f"docs/specs/SPEC-{issue_number}*.md",
        f"docs/ux/UX-{issue_number}*.md",
        f"docs/reviews/REVIEW-{issue_number}*.md",
    ]

    for pattern in patterns:
        docs.extend(repo_root.glob(pattern))

    return sorted(docs)


def _convert_md_to_docx(md_path: Path, docx_path: Path, _template: Optional[Path]) -> None:
    """Convert markdown to DOCX using python-docx directly."""
    from markdown_it import MarkdownIt

    # Parse markdown content
    md = MarkdownIt("commonmark", {"breaks": True, "html": False})
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    tokens = md.parse(content)

    # Create Word document directly
    _create_word_document_direct(tokens, docx_path, md_path.stem)


def _create_word_document_direct(tokens: List, output_path: Path,
                                  document_title: str) -> None:
    """Create Word document using python-docx directly."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX export.\n"
            "Install with: pip install python-docx"
        ) from exc

    logger.info("Creating Word document: %s", output_path)

    # Create new document
    doc = Document()

    # Set document properties
    core_properties = doc.core_properties
    core_properties.title = document_title
    core_properties.author = "ContextWeave Export"

    # Process markdown tokens and add content
    i = 0
    while i < len(tokens):
        token = tokens[i]

        try:
            if token.type == "heading_open":
                # Get heading level and text
                level = int(token.tag[1])
                i += 1
                if i < len(tokens) and tokens[i].type == "inline":
                    text = tokens[i].content
                    paragraph = doc.add_heading(text, level=level)
                    logger.debug("Added heading (level %d): %s...", level, text[:50])

            elif token.type == "paragraph_open":
                # Get paragraph text
                i += 1
                if i < len(tokens) and tokens[i].type == "inline":
                    text = tokens[i].content
                    if text:  # Only add non-empty paragraphs
                        doc.add_paragraph(text)
                        logger.debug("Added paragraph: %s...", text[:50])

            elif token.type == "code_block" or token.type == "fence":
                # Add code block as styled paragraph
                code_text = token.content
                if code_text:
                    paragraph = doc.add_paragraph(code_text)
                    # Style as code
                    for run in paragraph.runs:
                        run.font.name = "Courier New"
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    logger.debug("Added code block: %d chars", len(code_text))

            elif token.type == "table_open":
                # Handle tables - collect table data
                table_data = _extract_table_from_tokens(tokens, i)
                if table_data and isinstance(table_data, list) and len(table_data) > 0:
                    rows = len(table_data)
                    cols = len(table_data[0])

                    # Create table
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Light Grid Accent 1'

                    # Fill table data
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < cols:
                                table.rows[row_idx].cells[col_idx].text = cell_data

                    logger.debug("Added table: %dx%d", rows, cols)

            elif token.type == "hr":
                # Add horizontal rule as page break
                doc.add_page_break()
                logger.debug("Added page break")

        except (ValueError, KeyError, IndexError) as e:
            logger.warning("Failed to add element (token type: %s): %s", token.type, e)

        i += 1

    # Save document
    doc.save(str(output_path))
    logger.info("Document created successfully: %s", output_path)


def _convert_md_to_pdf(md_path: Path, pdf_path: Path) -> None:
    """Convert markdown to PDF using pandoc.

    Note: Direct MD->PDF requires pandoc to be installed.
    For Windows users, convert to DOCX first, then use docx2pdf.
    """
    try:
        subprocess.run([
            "pandoc",
            str(md_path),
            "-o", str(pdf_path),
            "--pdf-engine=xelatex",
            "-V", "geometry:margin=1in"
        ], check=True)
    except FileNotFoundError as exc:
        raise RuntimeError(
            "PDF conversion requires pandoc:\n"
            "  Install: https://pandoc.org/installing.html\n"
            "\n"
            "Alternative for Windows:\n"
            "  1. Export to DOCX: context-weave export document file.md --format docx\n"
            "  2. Install docx2pdf: pip install docx2pdf\n"
            "  3. Convert: python -m docx2pdf file.docx"
        ) from exc


def _convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """Convert DOCX to PDF using platform-specific tools.

    Args:
        docx_path: Path to input DOCX file
        pdf_path: Path to output PDF file

    Raises:
        RuntimeError: If PDF conversion fails
    """
    logger.info("Converting DOCX to PDF: %s -> %s", docx_path, pdf_path)

    import platform

    try:
        if platform.system() == "Windows":
            # Try docx2pdf (uses Microsoft Word)
            try:
                from docx2pdf import convert
                convert(str(docx_path), str(pdf_path))
                logger.info("PDF created using docx2pdf (Microsoft Word)")
                return
            except ImportError:
                logger.warning("docx2pdf not available, using weasyprint fallback")

            # No fallback available - require docx2pdf
            raise RuntimeError(
                "Windows PDF conversion requires docx2pdf:\n"
                "Install with: pip install docx2pdf\n"
                "Note: Requires Microsoft Word to be installed"
            )
        else:
            # Try LibreOffice on Linux/Mac
            subprocess.run([
                "soffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(pdf_path.parent),
                str(docx_path)
            ], check=True)
            logger.info("PDF created using LibreOffice")

    except (subprocess.CalledProcessError, FileNotFoundError, OSError) as e:
        raise RuntimeError(
            f"PDF conversion failed: {e}\n"
            "Install one of:\n"
            "  1. pip install docx2pdf (Windows, requires MS Word)\n"
            "  2. sudo apt install libreoffice (Linux)\n"
            "  3. brew install libreoffice (Mac)"
        ) from e

